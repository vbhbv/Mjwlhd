# -*- coding: utf-8 -*-
"""
Academic Knowledge Telegram Bot - Stable Production Version
"""
import os
import re
import uuid
import time
import logging
import sqlite3
import xml.etree.ElementTree as ET
from contextlib import contextmanager
from docx import Document
from docx.shared import RGBColor, Pt
from deep_translator import GoogleTranslator
from telegram import Update
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackContext

# 1. Logging Setup
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[logging.StreamHandler(), logging.FileHandler("bot_production.log", encoding="utf-8")]
)

DB_NAME = "bot_academic_memory.db"
VAULT_DIR = "bot_vault"
os.makedirs(VAULT_DIR, exist_ok=True)

# 2. Database Layer (SQLite + WAL Mode)
def init_db():
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("PRAGMA journal_mode=WAL;")
        conn.execute("PRAGMA synchronous=NORMAL;")
        cursor = conn.cursor()
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS translation_memory (
                text_hash TEXT PRIMARY KEY,
                original_text TEXT,
                translated_text TEXT,
                domain TEXT
            )
        """)
        conn.commit()

@contextmanager
def get_db():
    conn = sqlite3.connect(DB_NAME, timeout=30.0)
    conn.execute("PRAGMA journal_mode=WAL;")
    try:
        yield conn
    finally:
        conn.close()

class LocalTranslationMemory:
    @staticmethod
    def get_hash(text: str, domain: str) -> str:
        import hashlib
        combined = f"{domain.strip().lower()}:{text.strip()}"
        return hashlib.sha256(combined.encode('utf-8')).hexdigest()

    @classmethod
    def lookup(cls, text: str, domain: str) -> str:
        text_hash = cls.get_hash(text, domain)
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT translated_text FROM translation_memory WHERE text_hash = ?", (text_hash,))
            row = cursor.fetchone()
            return row[0] if row else None

    @classmethod
    def save(cls, text: str, translated: str, domain: str):
        text_hash = cls.get_hash(text, domain)
        with get_db() as conn:
            try:
                conn.execute("INSERT OR IGNORE INTO translation_memory VALUES (?, ?, ?, ?)", (text_hash, text, translated, domain))
                conn.commit()
            except Exception as e:
                logging.error(f"Error saving to cache: {e}")

# 3. Resource and Guardrail Validator
class StrictResourceValidator:
    MAX_PARAGRAPHS = 1500  
    MAX_TABLES = 30

    @classmethod
    def validate(cls, file_path: str):
        import zipfile
        if not zipfile.is_zipfile(file_path):
            raise ValueError("الملف غير صالح أو تالف.")
        with zipfile.ZipFile(file_path, 'r') as zf:
            with zf.open("word/document.xml") as doc_xml:
                tree = ET.parse(doc_xml)
                root = tree.getroot()
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                p_count = len(root.findall('.//w:p', namespaces))
                t_count = len(root.findall('.//w:tbl', namespaces))
                
                if p_count > cls.MAX_PARAGRAPHS or t_count > cls.MAX_TABLES:
                    raise ValueError(f"الملف ضخم جداً حوسبياً لدعم خادمك المحدود! (الفقرات: {p_count}/{cls.MAX_PARAGRAPHS})")

# 4. Document Object Model (DOM) Processor
class NodeType:
    PARAGRAPH = "paragraph"
    TABLE_CELL = "table_cell"

class DocumentNode:
    def __init__(self, node_id: int, node_type: str, raw_object, text: str):
        self.node_id = node_id
        self.node_type = node_type
        self.raw_object = raw_object
        self.original_text = text
        self.processed_text = text
        self.translated_text = None

class DocumentDOMProcessor:
    def __init__(self, file_path: str):
        self.doc = Document(file_path)
        self.nodes = []
        self._build_dom()

    def _build_dom(self):
        node_id = 0
        for p in self.doc.paragraphs:
            if p.text.strip():
                self.nodes.append(DocumentNode(node_id, NodeType.PARAGRAPH, p, p.text))
                node_id += 1
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for cell_p in cell.paragraphs:
                        if cell_p.text.strip():
                            self.nodes.append(DocumentNode(node_id, NodeType.TABLE_CELL, cell_p, cell_p.text))
                            node_id += 1

    def save_and_format(self, output_path: str):
        for node in self.nodes:
            if node.translated_text:
                try:
                    run = node.raw_object.add_run(f"\n{node.translated_text}")
                    run.font.color.rgb = RGBColor(28, 116, 182) 
                    run.font.size = Pt(10)
                    run.italic = True
                except Exception:
                    pass
        self.doc.save(output_path)

# 5. Core Pipeline Engine
class ProcessingEngine:
    @staticmethod
    def process_document(dom: DocumentDOMProcessor, domain: str) -> dict:
        tokens_map = {}
        token_id = 1
        
        patterns = [
            r'\b[OΩΘo]\s*\(\s*[nlog\d\s\+\-\*\/²\^]+\s*\)',
            r'\b[A-Za-z]\s*\(\s*[A-Za-z0-9_\s|]+\s*\)',
            r'\b(?:IEEE\s+802\.\d+[a-z]*|RFC\s+\d+|GPT-[\d\.]+)\b'
        ]
        
        for node in dom.nodes:
            text = node.processed_text
            for pattern in patterns:
                for match in set(re.findall(pattern, text)):
                    placeholder = f" <PROTECTEDENTITY id='{token_id}'> "
                    tokens_map[placeholder.strip()] = match
                    text = text.replace(match, placeholder)
                    token_id += 1
            node.processed_text = text

        translator = GoogleTranslator(source='en', target='ar')
        chunk_size = 10  
        
        for i in range(0, len(dom.nodes), chunk_size):
            chunk = dom.nodes[i:i+chunk_size]
            xml_packet = "<ROOT>"
            for node in chunk:
                cached = LocalTranslationMemory.lookup(node.processed_text, domain)
                if cached:
                    node.translated_text = cached
                else:
                    xml_packet += f"<SEG id='{node.node_id}'>{node.processed_text}</SEG>"
            xml_packet += "</ROOT>"

            if xml_packet != "<ROOT></ROOT>":
                try:
                    translated_packet = translator.translate(xml_packet)
                    root = ET.fromstring(f"<RESULT>{translated_packet}</RESULT>")
                    for seg in root.findall('.//seg'):
                        n_id = int(seg.get('id'))
                        trans_text = seg.text if seg.text else ""
                        
                        for placeholder, original_val in tokens_map.items():
                            trans_text = trans_text.replace(placeholder, original_val)
                            trans_text = trans_text.replace(placeholder.strip(), original_val)
                        
                        for node in chunk:
                            if node.node_id == n_id:
                                node.translated_text = trans_text
                                LocalTranslationMemory.save(node.original_text, trans_text, domain)
                except Exception as e:
                    logging.error(f"Batch translation error: {e}")
                    time.sleep(1)

        full_text_en = " ".join([n.original_text for n in dom.nodes])
        acronyms = set(re.findall(r'\b[A-Z]{3,4}\b', full_text_en))
        
        flashcards = []
        for node in dom.nodes:
            if "defined as" in node.original_text.lower() or "is the process of" in node.original_text.lower():
                try:
                    concept = node.original_text.split('is')[0].strip()
                    flashcards.append(f"• {concept} ➔ {node.translated_text}")
                except Exception:
                    pass

        return {
            "acronyms": list(acronyms)[:10],
            "flashcards": flashcards[:5]
        }

# 6. Telegram Bot Handlers
def start(update: Update, context: CallbackContext):
    update.message.reply_text(
        "👋 مرحباً بك في منصة معالجة المعرفة الأكاديمية السيادية!\n\n"
        "ℹ️ هذا البوت مصمم لتوفير استهلاك السيرفر بالكامل وحماية الموارد.\n"
        "📂 أرسل لي أي ملف مستندات (Word .docx) مكتوب بالإنجليزية، وسأقوم بترجمته صياغياً وبنائياً مع تلوين مزدوج واستخراج بطاقات المراجعة الذكية آلياً."
    )

def handle_document(update: Update, context: CallbackContext):
    doc_file = update.message.document
    if not doc_file.file_name.endswith('.docx'):
        update.message.reply_text("❌ عذراً، البوت يدعم ملفات Word بصيغة (.docx) فقط حالياً.")
        return

    status_msg = update.message.reply_text("⏳ جاري تأمين المستند وفحص القيود الحوسبية في السيرفر...")
    
    job_id = str(uuid.uuid4())
    input_path = os.path.join(VAULT_DIR, f"{job_id}_in.docx")
    output_path = os.path.join(VAULT_DIR, f"{job_id}_out.docx")

    try:
        tg_file = context.bot.get_file(doc_file.file_id)
        tg_file.download(custom_path=input_path)

        StrictResourceValidator.validate(input_path)
        
        status_msg.edit_text("🧠 يجري الآن قراءة هيكل الـ DOM للمستند وعزل المعادلات والرموز العلمية...")
        dom = DocumentDOMProcessor(input_path)
        
        status_msg.edit_text("🚀 جاري معالجة دفعات الترجمة الذكية عبر الكاش المحلي المتزامن...")
        knowledge_data = ProcessingEngine.process_document(dom, "academic_general")
        
        status_msg.edit_text("💾 جاري إعادة بناء ملف الوورد النهائي وحقن التلوين المزدوج...")
        dom.save_and_format(output_path)

        status_msg.edit_text("✅ اكتملت المعالجة الهندسية بنجاح! جاري رفع ملفك...")
        with open(output_path, 'rb') as f:
            update.message.reply_document(document=f, filename=f"Bilingual_Result_{doc_file.file_name}")

        knowledge_report = "🎓 **حزمة المراجعة المعرفية المستخرجة تلقائياً:**\n\n"
        if knowledge_data["acronyms"]:
            knowledge_report += f"🔍 **أهم الاختصارات الواردة:** {', '.join(knowledge_data['acronyms'])}\n\n"
        if knowledge_data["flashcards"]:
            knowledge_report += "🗂 **بطاقات استذكار مقترحة (Flashcards):**\n" + "\n".join(knowledge_data["flashcards"])
        else:
            knowledge_report += "📝 لم يتم رصد جمل تعريفية صريحة لصناعة بطاقات استذكار تلقائية."

        update.message.reply_text(knowledge_report, parse_mode="Markdown")

    except Exception as e:
        logging.error(f"Error handling job {job_id}: {e}", exc_info=True)
        update.message.reply_text(f"❌ حدث خطأ أثناء المعالجة: {str(e)}")
    finally:
        if os.path.exists(input_path): 
            os.remove(input_path)
        if os.path.exists(output_path): 
            os.remove(output_path)
        if 'status_msg' in locals():
            try:
                status_msg.delete()
            except Exception:
                pass

def main():
    init_db()
    TOKEN = "6807502954:AAH5tOwXCjRXtF65wQFEDSkYeFBYIgUjblg"
    
    updater = Updater(TOKEN, use_context=True)
    dp = updater.dispatcher

    dp.add_handler(CommandHandler("start", start))
    dp.add_handler(MessageHandler(Filters.document, handle_document))

    logging.info("🚀 البوت الأكاديمي يعمل الآن بنظام التقشف الحوسبي الآمن...")
    updater.start_polling()
    updater.idle()

if __name__ == '__main__':
    main()
